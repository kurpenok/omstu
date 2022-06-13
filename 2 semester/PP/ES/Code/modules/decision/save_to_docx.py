import os

from docx import Document


def save(self) -> None:
    wd = os.getcwd()

    if os.path.exists(f"{wd}/Report.docx"):
        return

    document = Document()

    document.add_heading("Report", 0)
    document.add_paragraph(f"{self.surname} {self.name}: {self.answer}")
    document.save("Report.docx")

