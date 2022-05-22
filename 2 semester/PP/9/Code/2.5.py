from docx import Document

document = Document()

document.add_heading("Blood test", 0)

records = [
        "WBC",
        "RBC",
        "HGB",
        "HTC",
        "MCV",
        "MCH",
        "MCHC",
        "RDW",
        "PLT",
        "MPV",
        "PTC",
        "NEU",
        "LYMP",
        "MONO",
        "EFO",
        "Baso"
]

lenght = len(records)

table = document.add_table(rows=1, cols=3)

header = table.rows[0].cells
header[0].text = "indicator"
header[1].text = "noem"
header[2].text = "value"

for i in range(lenght):
    rows = table.add_row().cells
    rows[0].text = records[i]

document.save("analysis.docx")

