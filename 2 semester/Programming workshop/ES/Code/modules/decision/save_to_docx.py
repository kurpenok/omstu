import os

from docx import Document


def save(self) -> None:
    wd = os.getcwd()

    if not os.path.exists(f"{wd}/reports/"):
        os.mkdir(f"{wd}/reports/")


    if os.path.exists(f"{wd}/reports/Report_{self.surname}_{self.name}.docx"):
        return

    document = Document()

    document.add_heading("Report", 0)
    document.add_paragraph(f"{self.surname} {self.name} {self.patronymic}: {self.answer}")
    document.save(f"reports/Report_{self.surname}_{self.name}.docx")

